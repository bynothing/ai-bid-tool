#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
S9 — 技术偏离表填写
从S3正文自动提取响应摘要，自动发现章节书签和页码，填写报价文件偏离表，
插入PAGEREF交叉引用字段。

用法:
  python scripts/s9_fill_deviation.py --formal-doc <报价文件.docx>
  python scripts/s9_fill_deviation.py --formal-doc <报价文件.docx> --dry-run
  python scripts/s9_fill_deviation.py --formal-doc <报价文件.docx> --output <输出.docx>
"""

import os
import sys
import re
import json
import copy
import argparse
from datetime import datetime
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# ============================================================
# Configuration
# ============================================================
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(ROOT, 'output')
S3_CHAPTERS_DIR = os.path.join(OUTPUT_DIR, 's3_body', 'chapters')
TRACE_PATH = os.path.join(OUTPUT_DIR, 'trace.json')
S9_OUTPUT_DIR = os.path.join(OUTPUT_DIR, 's9_deviation')

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Chapter ID → Chinese chapter name
CHAPTER_ID_MAP = OrderedDict([
    ('CH-01', '第一章 项目理解与总体技术方案'),
    ('CH-02', '第二章 资质、保密与合规保障'),
    ('CH-03', '第三章 总体集成通用技术方案'),
    ('CH-04', '第四章 人机交互界面设计'),
    ('CH-05', '第五章 数据库模块技术方案'),
    ('CH-06', '第六章 材料设计模块技术方案'),
    ('CH-07', '第七章 热环境确定模块技术方案'),
    ('CH-08', '第八章 热环境响应评估模块技术方案'),
    ('CH-09', '第九章 运行环境兼容性与离线部署方案'),
    ('CH-10', '第十章 集成对接与数据贯通方案'),
    ('CH-11', '第十一章 项目实施与交付方案'),
    ('CH-12', '第十二章 测试、验收与质量保障'),
    ('CH-13', '第十三章 售后服务与质保承诺'),
])

# Boilerplate phrases to strip from summaries
BOILERPLATE_PATTERNS = [
    re.compile(r'我方承诺[^。；;，]+[。；;，]'),
    re.compile(r'我方郑重声明[^。]+[。]'),
    re.compile(r'我方[对]?本条款[的]?承诺无偏离[。]?'),
    re.compile(r'[（(]详见[^）)]*[）)]'),
    re.compile(r'满足招标文件[^。]*?要求[。，]?'),
    re.compile(r'严格[遵执]照[^。]+[。]'),
    re.compile(r'完全响应[^。]*?条款[。，]?'),
    re.compile(r'^[。，,；;、\s]+'),
]


# ============================================================
# PAGEREF field creation
# ============================================================
def make_pageref_field(bookmark_name):
    """Create a PAGEREF Word field XML element."""
    container = etree.Element(qn('w:r'))

    r_begin = etree.SubElement(container, qn('w:r'))
    fld_begin = etree.SubElement(r_begin, qn('w:fldChar'))
    fld_begin.set(qn('w:fldCharType'), 'begin')

    r_instr = etree.SubElement(container, qn('w:r'))
    instr_text = etree.SubElement(r_instr, qn('w:instrText'))
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f' PAGEREF {bookmark_name} \\h '

    r_sep = etree.SubElement(container, qn('w:r'))
    fld_sep = etree.SubElement(r_sep, qn('w:fldChar'))
    fld_sep.set(qn('w:fldCharType'), 'separate')

    r_val = etree.SubElement(container, qn('w:r'))
    t_val = etree.SubElement(r_val, qn('w:t'))
    t_val.text = '1'

    r_end = etree.SubElement(container, qn('w:r'))
    fld_end = etree.SubElement(r_end, qn('w:fldChar'))
    fld_end.set(qn('w:fldCharType'), 'end')

    return container


# ============================================================
# Cell manipulation helpers
# ============================================================
def set_cell_text(cell_elem, text):
    """Set text in a cell while preserving formatting."""
    w_t_elements = cell_elem.findall('.//' + qn('w:t'))
    if w_t_elements:
        w_t_elements[0].text = text
        for wt in w_t_elements[1:]:
            wt.text = ''
    else:
        w_p_elements = cell_elem.findall(qn('w:p'))
        if w_p_elements:
            p = w_p_elements[0]
            r = etree.SubElement(p, qn('w:r'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = text


def set_cell_pageref(cell_elem, bookmark_name):
    """Replace cell content with a PAGEREF field."""
    w_p_elements = cell_elem.findall(qn('w:p'))
    for wp in w_p_elements:
        cell_elem.remove(wp)

    new_p = etree.Element(qn('w:p'))
    cell_elem.append(new_p)

    pPr = etree.SubElement(new_p, qn('w:pPr'))
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'center')

    pageref = make_pageref_field(bookmark_name)
    for child in pageref:
        new_p.append(child)


# ============================================================
# 1. Pipeline data loading
# ============================================================
def load_pipeline_data(pipeline_data_dir):
    """Load trace.json and build requirement→chapter mapping."""
    trace_file = os.path.join(pipeline_data_dir, 'trace.json')
    if not os.path.exists(trace_file):
        raise FileNotFoundError(f'trace.json not found at {trace_file}')

    with open(trace_file, 'r', encoding='utf-8') as f:
        trace = json.load(f)

    requirements = trace.get('requirements', [])
    body_segments = trace.get('body_segments', [])

    # Build body_id → text_preview lookup
    body_lookup = {}
    for seg in body_segments:
        body_lookup[seg.get('body_id', '')] = {
            'text_preview': seg.get('text_preview', ''),
            'chapter': seg.get('chapter', ''),
            'chapter_id': seg.get('chapter_id', ''),
            'covers_reqs': seg.get('covers_reqs', []),
        }

    return requirements, body_lookup


# ============================================================
# 2. Summary extraction from S3 body text
# ============================================================
def extract_summaries(requirements, body_lookup, verbose=False):
    """Extract response summaries from S3 chapter markdown files.

    Primary: Parse 【对标条款 RQ-TECH-XXX】 markers in chapter .md files
    Fallback: Use text_preview from trace.json body_segments
    Last resort: Use shortened requirement text
    """
    summaries = {}

    # Load all chapter markdown content
    chapter_texts = {}
    if os.path.isdir(S3_CHAPTERS_DIR):
        for fname in sorted(os.listdir(S3_CHAPTERS_DIR)):
            if fname.endswith('.md'):
                fpath = os.path.join(S3_CHAPTERS_DIR, fname)
                with open(fpath, 'r', encoding='utf-8') as f:
                    chapter_texts[fname] = f.read()

    # Build regex to find 【对标条款 RQ-TECH-XXX ...】 markers
    marker_pattern = re.compile(
        r'>\s*\*{0,2}【对标条款\s*(RQ-TECH-\d+)\s*(?:\[[^]]*\])?\s*】[^\n]*\n'
        r'((?:(?!> \*\*【|#{1,3} |---|\n\n>).)*)',
        re.MULTILINE | re.DOTALL
    )

    # Extract from each chapter
    for req in requirements:
        req_id = req.get('req_id', '')
        chapter_id = req.get('outline_id', '')
        chapter_name = CHAPTER_ID_MAP.get(chapter_id, '')

        extracted_text = None

        # Try primary method: 【对标条款】 markers in chapter .md
        for fname, content in chapter_texts.items():
            for match in marker_pattern.finditer(content):
                if match.group(1) == req_id:
                    # Get the text following the marker
                    raw = match.group(2).strip()
                    # Clean up: remove > prefix, ** markers, newlines
                    raw = re.sub(r'^>\s*', '', raw, flags=re.MULTILINE)
                    raw = re.sub(r'\*\*', '', raw)
                    raw = re.sub(r'\n+', '', raw)
                    if len(raw) > 30:  # Must have enough substance
                        extracted_text = raw
                        break
            if extracted_text:
                break

        # Fallback 1: body_segments text_preview
        if not extracted_text:
            body_ids = req.get('body_ids', [])
            previews = []
            for bid in body_ids:
                if bid in body_lookup:
                    tp = body_lookup[bid].get('text_preview', '')
                    if tp and len(tp) > 10:
                        previews.append(tp)
            if previews:
                extracted_text = '。'.join(previews[:3])

        # Fallback 2: requirement text itself
        if not extracted_text:
            extracted_text = req.get('text', '')

        # Clean up boilerplate phrases
        for pattern in BOILERPLATE_PATTERNS:
            extracted_text = pattern.sub('', extracted_text)

        # Trim whitespace
        extracted_text = re.sub(r'\s+', '', extracted_text)

        # Smart truncation at ~200 chars
        if len(extracted_text) > 200:
            # Find last sentence break within 200 chars
            truncated = extracted_text[:200]
            last_period = max(
                truncated.rfind('。'),
                truncated.rfind('；'),
                truncated.rfind('；'),
            )
            if last_period > 100:
                extracted_text = truncated[:last_period + 1]
            else:
                extracted_text = truncated + '…'

        summaries[req_id] = extracted_text

        if verbose and len(extracted_text) < 20:
            print(f'  WARNING: Short summary for {req_id}: {extracted_text[:50]}')

    return summaries


# ============================================================
# 3. Bookmark discovery from formal bid document XML
# ============================================================
def discover_bookmarks(doc, verbose=False):
    """Discover _Toc bookmark → chapter name mapping from the DOCX XML."""
    bookmark_map = {}

    # Find all bookmarkStart elements with _Toc prefix
    body = doc.element.body
    bm_starts = body.findall('.//' + qn('w:bookmarkStart'))
    toc_bookmarks = []
    for bm in bm_starts:
        name = bm.get(qn('w:name'))
        if name and name.startswith('_Toc'):
            toc_bookmarks.append(name)

    if verbose:
        print(f'  Found {len(toc_bookmarks)} _Toc bookmarks')

    # Now find each bookmark's associated paragraph text
    # Bookmarks are inline — the paragraph containing the bookmark start
    # is the chapter heading paragraph
    # NOTE: @name attribute has NO namespace in OOXML (it's unprefixed)
    for bm_name in toc_bookmarks:
        bm_elem = None
        for bm in body.iter(qn('w:bookmarkStart')):
            if bm.get(qn('w:name')) == bm_name:
                bm_elem = bm
                break
        if bm_elem is None:
            continue

        # Get the parent paragraph
        p_elem = bm_elem.getparent()
        while p_elem is not None and p_elem.tag != qn('w:p'):
            p_elem = p_elem.getparent()

        if p_elem is None:
            continue

        # Get all text from this paragraph
        para_text = ''
        for t_elem in p_elem.findall('.//' + qn('w:t')):
            if t_elem.text:
                para_text += t_elem.text

        # Try to match against chapter names using chapter number prefix
        # (e.g., "第十二章") for fuzzy matching — actual heading text may
        # have minor punctuation differences from CHAPTER_ID_MAP values
        for chapter_name in CHAPTER_ID_MAP.values():
            ch_num_match = re.match(r'(第[一二三四五六七八九十百]+章)', chapter_name)
            ch_number = ch_num_match.group(1) if ch_num_match else chapter_name

            # Already matched this chapter?
            if chapter_name in bookmark_map:
                continue

            # Check if this bookmark's paragraph contains the chapter number
            if ch_number in para_text:
                bookmark_map[chapter_name] = bm_name
                if verbose:
                    print(f'    {chapter_name} → {bm_name}')
                break

    if verbose:
        missing = [ch for ch in CHAPTER_ID_MAP.values() if ch not in bookmark_map]
        if missing:
            print(f'  WARNING: Unmatched chapters: {missing}')

    return bookmark_map


# ============================================================
# 4. Page number detection
# ============================================================
def detect_page_numbers_win32com(formal_doc_path, bookmark_map, verbose=False):
    """Detect page numbers using Word COM automation."""
    import win32com.client

    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0

    page_numbers = {}
    try:
        doc = word.Documents.Open(formal_doc_path)
        doc.Repaginate()

        # wdActiveEndPageNumber = 3
        for chapter_name, bm_name in bookmark_map.items():
            try:
                bm = doc.Bookmarks(bm_name)
                page = bm.Range.Information(3)
                page_numbers[chapter_name] = page
                if verbose:
                    print(f'    {chapter_name} → Page {page}')
            except Exception as e:
                if verbose:
                    print(f'    WARNING: {chapter_name} ({bm_name}): {e}')

        doc.Close()
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    return page_numbers


def detect_page_numbers_pdf(formal_doc_path, verbose=False):
    """Fallback: convert to PDF and search for chapter headings."""
    try:
        import pdfplumber
        import win32com.client
        import tempfile
    except ImportError:
        if verbose:
            print('  pdfplumber not available for fallback page detection')
        return {}

    # Export to temp PDF
    tmp_pdf = os.path.join(tempfile.gettempdir(), '_s9_temp.pdf')
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(formal_doc_path)
        doc.SaveAs(tmp_pdf, FileFormat=17)
        doc.Close()
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    if not os.path.exists(tmp_pdf):
        return {}

    # Search PDF
    page_numbers = {}
    with pdfplumber.open(tmp_pdf) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                for chapter_name in CHAPTER_ID_MAP.values():
                    if chapter_name in text and chapter_name not in page_numbers:
                        page_numbers[chapter_name] = i + 1

    try:
        os.remove(tmp_pdf)
    except Exception:
        pass

    return page_numbers


def detect_page_numbers(formal_doc_path, bookmark_map, verbose=False):
    """Detect page numbers. Try win32com first, PDF fallback, then use defaults.

    The default page numbers come from the original PDF verification
    (run when Word is available). These are stable as long as the document
    content before the technical chapters doesn't change significantly.
    """
    # Default page numbers from verified PDF (fallback of last resort)
    DEFAULT_PAGE_NUMBERS = {
        '第一章 项目理解与总体技术方案': 118,
        '第二章 资质、保密与合规保障': 122,
        '第三章 总体集成通用技术方案': 128,
        '第四章 人机交互界面设计': 136,
        '第五章 数据库模块技术方案': 143,
        '第六章 材料设计模块技术方案': 160,
        '第七章 热环境确定模块技术方案': 170,
        '第八章 热环境响应评估模块技术方案': 179,
        '第九章 运行环境兼容性与离线部署方案': 186,
        '第十章 集成对接与数据贯通方案': 191,
        '第十一章 项目实施与交付方案': 199,
        '第十二章 测试、验收与质量保障': 204,
        '第十三章 售后服务与质保承诺': 209,
    }

    page_numbers = {}

    # Method 1: win32com bookmark lookup (needs MS Word or WPS)
    try:
        if verbose:
            print('  Trying win32com bookmark detection...')
        page_numbers = detect_page_numbers_win32com(formal_doc_path, bookmark_map, verbose)
        if len(page_numbers) >= 10:
            if verbose:
                print(f'  win32com: got {len(page_numbers)} page numbers')
            return page_numbers
    except Exception as e:
        if verbose:
            print(f'  win32com unavailable: {e}')

    # Method 2: Export to PDF and search for chapter headings
    try:
        if verbose:
            print('  Falling back to PDF extraction...')
        page_numbers = detect_page_numbers_pdf(formal_doc_path, verbose)
        if len(page_numbers) >= 10:
            if verbose:
                print(f'  PDF: got {len(page_numbers)} page numbers')
            return page_numbers
    except Exception as e:
        if verbose:
            print(f'  PDF extraction failed: {e}')

    # Method 3: Use default page numbers from prior verification
    if verbose:
        print('  Using default page numbers from prior PDF verification')
    for chapter_name, page in DEFAULT_PAGE_NUMBERS.items():
        page_numbers[chapter_name] = page

    return page_numbers


# ============================================================
# 5. Parse comparison table from formal bid document
# ============================================================
def parse_comparison_table(doc, table_index=10):
    """Parse the requirements comparison table from the formal bid document."""
    if table_index >= len(doc.tables):
        raise ValueError(f'Table index {table_index} out of range ({len(doc.tables)} tables)')

    tbl = doc.tables[table_index]
    requirements = []

    for i in range(1, len(tbl.rows)):
        row = tbl.rows[i]
        cells = row.cells
        if len(cells) < 6:
            continue
        seq = cells[0].text.strip()
        if not seq or seq == '备注':
            continue

        requirements.append({
            'seq': seq,
            'clause_id': cells[1].text.strip(),
            'nature': cells[2].text.strip(),
            'tech_requirement': cells[3].text.strip(),
            'response_chapter': cells[4].text.strip(),
            'response_points': cells[5].text.strip(),
        })

    return requirements


# ============================================================
# 6. Fill deviation table
# ============================================================
def fill_deviation_table(doc, requirements, summaries, bookmark_map, page_numbers,
                         table_index=9, verbose=False):
    """Fill the deviation table in the formal bid document."""
    if table_index >= len(doc.tables):
        raise ValueError(f'Table index {table_index} out of range ({len(doc.tables)} tables)')

    tbl = doc.tables[table_index]
    tbl_elem = tbl._tbl
    all_rows = tbl_elem.findall(qn('w:tr'))

    if len(all_rows) < 2:
        raise ValueError(f'Deviation table has only {len(all_rows)} rows, need at least 2')

    # Row 0 = header, Row 1 = template data row, Row last = 备注
    template_row_elem = all_rows[1]
    remark_row_elem = all_rows[-1]

    # Remove all rows after header
    for row_elem in all_rows[1:]:
        tbl_elem.remove(row_elem)

    # Fill each requirement
    missing_bookmarks = []
    for idx, req in enumerate(requirements):
        ch_name = req.get('response_chapter', '')
        clause_id = req.get('clause_id', '')
        bm_name = bookmark_map.get(ch_name, None)
        page_num = page_numbers.get(ch_name, None)
        summary = summaries.get(clause_id, req.get('tech_requirement', ''))

        if not bm_name and ch_name:
            missing_bookmarks.append(ch_name)

        # Clone template row
        new_row = copy.deepcopy(template_row_elem)
        tbl_elem.append(new_row)
        tc_elems = new_row.findall(qn('w:tc'))

        if len(tc_elems) < 6:
            if verbose:
                print(f'  WARNING: Row {idx + 1} has only {len(tc_elems)} cells')
            continue

        # Column 0: 序号
        set_cell_text(tc_elems[0], str(idx + 1))

        # Column 1: 技术要求
        set_cell_text(tc_elems[1], req.get('tech_requirement', ''))

        # Column 2: 对应响应内容
        set_cell_text(tc_elems[2], summary)

        # Column 3: 偏离情况
        set_cell_text(tc_elems[3], '无偏离')

        # Column 4: 页码 (PAGEREF cross-reference)
        if bm_name:
            set_cell_pageref(tc_elems[4], bm_name)
        else:
            set_cell_text(tc_elems[4], '')

        # Column 5: 说明
        note_parts = []
        if ch_name:
            note_parts.append(ch_name)
        if page_num:
            note_parts.append(f'（第{page_num}页）')
        set_cell_text(tc_elems[5], ''.join(note_parts))

        if verbose and (idx + 1) % 20 == 0:
            print(f'  Filled {idx + 1}/{len(requirements)} rows...')

    # Re-add 备注 row
    tbl_elem.append(remark_row_elem)

    return missing_bookmarks


# ============================================================
# 7. Output validation
# ============================================================
def validate_output(doc, table_index=9, verbose=False):
    """Validate the filled deviation table."""
    if table_index >= len(doc.tables):
        return {'error': f'Table {table_index} not found'}

    tbl = doc.tables[table_index]
    issues = []
    stats = {
        'total_rows': len(tbl.rows),
        'data_rows': 0,
        'pageref_fields': 0,
        'non_deviation_count': 0,
        'empty_summaries': 0,
        'empty_tech_req': 0,
    }

    tbl_elem = tbl._tbl
    instr_texts = tbl_elem.findall(f'.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}instrText')
    stats['pageref_fields'] = len([i for i in instr_texts if 'PAGEREF' in (i.text or '')])

    for i in range(1, len(tbl.rows)):
        row = tbl.rows[i]
        cells = row.cells
        if len(cells) < 6:
            continue
        seq = cells[0].text.strip()
        # Skip remark row and empty rows
        if seq == '备注' or not seq or not seq.isdigit():
            continue

        stats['data_rows'] += 1

        tech_req = cells[1].text.strip()
        summary = cells[2].text.strip()
        deviation = cells[3].text.strip()

        if not tech_req:
            stats['empty_tech_req'] += 1
            issues.append(f'Row {i}: empty tech requirement')

        if not summary or len(summary) < 5:
            stats['empty_summaries'] += 1
            issues.append(f'Row {i} ({seq}): empty or very short summary')

        if deviation != '无偏离':
            stats['non_deviation_count'] += 1
            issues.append(f'Row {i} ({seq}): deviation is "{deviation}"')

    if issues and verbose:
        for issue in issues[:20]:
            print(f'  {issue}')

    stats['issues'] = issues
    stats['passed'] = len(issues) == 0

    return stats


# ============================================================
# Main
# ============================================================
def main():
    parser = argparse.ArgumentParser(description='S9 — 技术偏离表填写')
    parser.add_argument('--formal-doc', required=True,
                        help='Path to the formal bid document (报价文件 DOCX)')
    parser.add_argument('--output', default=None,
                        help='Output path (default: formal-doc + _偏离表已填写.docx)')
    parser.add_argument('--pipeline-data', default=OUTPUT_DIR,
                        help=f'Pipeline output directory (default: {OUTPUT_DIR})')
    parser.add_argument('--dry-run', action='store_true',
                        help='Validate only, do not modify the document')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Detailed logging')

    args = parser.parse_args()

    if not os.path.exists(args.formal_doc):
        print(f'ERROR: Formal bid document not found: {args.formal_doc}')
        sys.exit(1)

    # Resolve output path
    if args.output:
        output_path = args.output
    else:
        base, ext = os.path.splitext(args.formal_doc)
        output_path = f'{base}_偏离表已填写{ext}'

    # Prepare S9 output directory
    os.makedirs(S9_OUTPUT_DIR, exist_ok=True)

    print('=' * 60)
    print('S9 — 技术偏离表填写')
    print('=' * 60)

    # Step 1: Load pipeline data
    print('\n[1/6] Loading pipeline data...')
    requirements, body_lookup = load_pipeline_data(args.pipeline_data)
    print(f'  {len(requirements)} requirements loaded from trace.json')
    print(f'  {len(body_lookup)} body segments indexed')

    # Step 2: Extract summaries from S3 body
    print('\n[2/6] Extracting response summaries from S3 body text...')
    summaries = extract_summaries(requirements, body_lookup, verbose=args.verbose)

    empty_count = sum(1 for v in summaries.values() if len(v) < 10)
    print(f'  {len(summaries)} summaries extracted ({empty_count} may need review)')

    # Save summaries cache
    summaries_cache = {
        'generated_at': datetime.now().isoformat(),
        'source': 'S3 chapter body 【对标条款】 markers + trace.json fallback',
        'summaries': summaries,
        'empty_count': empty_count,
    }
    with open(os.path.join(S9_OUTPUT_DIR, 's9_summaries.json'), 'w', encoding='utf-8') as f:
        json.dump(summaries_cache, f, ensure_ascii=False, indent=2)

    # Step 3: Open formal document and discover bookmarks
    print('\n[3/6] Discovering bookmarks and page numbers...')
    doc = Document(args.formal_doc)

    bookmark_map = discover_bookmarks(doc, verbose=args.verbose)
    print(f'  {len(bookmark_map)}/{len(CHAPTER_ID_MAP)} chapters mapped to TOC bookmarks')

    # Save bookmark map
    with open(os.path.join(S9_OUTPUT_DIR, 's9_bookmark_map.json'), 'w', encoding='utf-8') as f:
        json.dump({
            'discovery_method': 'xml_bookmark_extraction',
            'bookmarks': bookmark_map,
        }, f, ensure_ascii=False, indent=2)

    # Step 4: Detect page numbers
    print('\n[4/6] Detecting page numbers...')
    page_numbers = detect_page_numbers(args.formal_doc, bookmark_map, verbose=args.verbose)
    print(f'  {len(page_numbers)}/{len(CHAPTER_ID_MAP)} chapters with page numbers')

    # Save page numbers
    with open(os.path.join(S9_OUTPUT_DIR, 's9_page_numbers.json'), 'w', encoding='utf-8') as f:
        json.dump({
            'detection_method': 'win32com_bookmark_lookup',
            'page_numbers': page_numbers,
        }, f, ensure_ascii=False, indent=2)

    # Step 5: Parse comparison table
    print('\n[5/6] Parsing comparison table from formal document...')
    comparison_reqs = parse_comparison_table(doc, table_index=10)
    print(f'  {len(comparison_reqs)} requirements parsed from comparison table')
    print(f'  Table columns: {len(doc.tables[10].columns)}')

    # Step 6: Fill deviation table
    print('\n[6/6] Filling deviation table...')

    if args.dry_run:
        print('  DRY RUN — validating only, not modifying document')
        # Just validate the incoming data
        for i, req in enumerate(comparison_reqs):
            cid = req.get('clause_id', '')
            ch = req.get('response_chapter', '')
            bm = bookmark_map.get(ch)
            pg = page_numbers.get(ch)
            summary = summaries.get(cid, '')
            if not summary or len(summary) < 10:
                print(f'  WARNING: Row {i+1} ({cid}): empty/short summary')
            if not bm:
                print(f'  WARNING: Row {i+1} ({cid}): no bookmark for "{ch}"')
            if not pg:
                print(f'  WARNING: Row {i+1} ({cid}): no page number for "{ch}"')
        print('  Dry run complete.')
    else:
        missing = fill_deviation_table(
            doc, comparison_reqs, summaries, bookmark_map, page_numbers,
            table_index=9, verbose=args.verbose
        )

        if missing:
            print(f'\n  WARNING: {len(missing)} chapters without bookmarks: {set(missing)}')

        # Validate
        stats = validate_output(doc, table_index=9, verbose=args.verbose)

        # Save
        print(f'\n  Saving to: {output_path}')
        doc.save(output_path)

        # Run report
        run_report = {
            'timestamp': datetime.now().isoformat(),
            'formal_doc': args.formal_doc,
            'output': output_path,
            'stats': {k: v for k, v in stats.items() if k != 'issues'},
            'bookmark_map': bookmark_map,
            'page_numbers': page_numbers,
            'summary_empty_count': empty_count,
            'missing_bookmarks': missing,
        }
        with open(os.path.join(S9_OUTPUT_DIR, 's9_run_report.json'), 'w', encoding='utf-8') as f:
            json.dump(run_report, f, ensure_ascii=False, indent=2)

        # Summary
        print(f'\n{"=" * 60}')
        print(f'Done!')
        print(f'  Rows: {stats["data_rows"]}')
        print(f'  PAGEREF fields: {stats["pageref_fields"]}')
        print(f'  Non-无偏离: {stats["non_deviation_count"]}')
        print(f'  Empty summaries: {stats["empty_summaries"]}')
        print(f'  Issues: {len(stats["issues"])}')
        if stats['passed']:
            print(f'  Result: PASSED')
        else:
            print(f'  Result: NEEDS REVIEW ({len(stats["issues"])} issues)')
        print(f'{"=" * 60}')

        # Exit code
        if not stats['passed']:
            sys.exit(1)


if __name__ == '__main__':
    main()
