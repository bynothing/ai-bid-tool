"""Stage 8: Export synthesized markdown to Word (.docx).

Converts s6_combined_bid_document.md to a properly formatted .docx file with:
- Heading styles (Heading 1/2/3 for chapter/section/subsection)
- Embedded images with captions
- Bold, italic, code inline formatting
- Blockquote-style requirement annotations
- Page breaks between chapters
- Markdown table → Word table conversion
- Bidding requirements comparison appendix
- Consistent Chinese document formatting
"""
import os
import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Inline parsing patterns ───
BOLD_PATTERN = re.compile(r'\*\*(.+?)\*\*')
CODE_PATTERN = re.compile(r'`([^`]+)`')
IMAGE_PATTERN = re.compile(r'^!\[(.*?)\]\((.*?)\)$')
HEADING_PATTERN = re.compile(r'^(#{1,3})\s+(.+)$')
BLOCKQUOTE_PATTERN = re.compile(r'^>\s+(.+)$')
HR_PATTERN = re.compile(r'^---\s*$')
LIST_ITEM_PATTERN = re.compile(r'^(-|\d+\.)\s+(.+)$')
CAPTION_PATTERN = re.compile(r'^\*(.+)\*$')
FENCE_PATTERN = re.compile(r'^```(\w*)\s*$')
TABLE_ROW_PATTERN = re.compile(r'^\|(.+)\|$')

# ─── Utility functions ───

# Regex patterns for RQ-TECH annotation removal
_RQ_MARKER = re.compile(r'【对标条款\s*RQ-TECH-\d+(?:\s*\[[^]]*\])?\s*】[，,，]?\s*')
_RQ_PAREN = re.compile(r'[（(]\s*RQ-TECH-\d+(?:\s*[、，,，]\s*(?:RQ-TECH-)?\d+)*\s*[）)]')
_RQ_PREFIX = re.compile(r'针对\s*RQ-TECH-\d+(?:[、，,]\s*RQ-TECH-\d+)*\s*要求[，,，]\s*')
_RQ_TRAILING = re.compile(
    r'[，,]\s*'
    r'(?:满足|响应|符合|遵守|覆盖|完全响应|满足并完全响应|理解并完全响应|保证)'
    r'(?:招标文件\s*)?'
    r'[^。]*?'
    r'RQ-TECH-\d+'
    r'[^。]*?'
    r'[。]'
)
_RQ_STANDALONE = re.compile(
    r'(?:^|[。])\s*'
    r'(?:我方(?:郑重声明，)?\s*)?'
    r'(?:我方\s*)?'
    r'(?:满足|响应|符合|遵守|覆盖|完全响应|满足并完全响应|理解并完全响应|保证)\s*'
    r'(?:招标文件\s*)?\s*'
    r'RQ-TECH-\d+'
    r'[^。]*?'
    r'[。]'
)
_RQ_ORPHAN_SENTENCE = re.compile(r'[。]\s*(?:本方案无偏离|响应无偏离|完全响应|全部满足|完全满足|完整覆盖|无偏离)[。]')
_RQ_BARE = re.compile(r'RQ-TECH-\d+(?:\s*[、，,，和]\s*(?:RQ-TECH-)?\d+)*')
_RQ_WS = re.compile(r'\s{2,}')
_RQ_WS_PUNCT = re.compile(r'\s+([，,。；;：、])')
_RQ_DOUBLE_PERIOD = re.compile(r'[。][。]')


def clean_body_text(text):
    """Remove all RQ-TECH requirement number references from body text.

    These are AI-generated annotations not present in the original tender document.
    The cleanup handles: inline 【对标条款】 markers, parenthetical (RQ-TECH-XXX) refs,
    trailing compliance clauses, standalone compliance sentences, and bare RQ-TECH numbers.
    """
    if not text or 'RQ-TECH-' not in text:
        return text

    # 1. Remove 【对标条款 RQ-TECH-XXX [label]】 inline markers
    text = _RQ_MARKER.sub('', text)

    # 2. Remove parenthetical RQ-TECH references: （RQ-TECH-XXX） or （RQ-TECH-XXX、XXX）
    text = _RQ_PAREN.sub('', text)

    # 3. Remove "针对 RQ-TECH-XXX 要求，" type prefixes
    text = _RQ_PREFIX.sub('', text)

    # 4. Remove trailing compliance clauses: "，满足 RQ-TECH-XXX ...。"
    text = _RQ_TRAILING.sub('。', text)

    # 5. Remove standalone compliance sentences
    text = _RQ_STANDALONE.sub('', text)

    # 6. Remove orphaned "本方案无偏离。" / "响应无偏离。" after RQ-TECH cleanup
    text = _RQ_ORPHAN_SENTENCE.sub('', text)

    # 7. Remove remaining bare RQ-TECH-\d+ references
    text = _RQ_BARE.sub('', text)

    # 8. Clean up whitespace and punctuation artifacts
    text = _RQ_WS.sub(' ', text)
    text = _RQ_WS_PUNCT.sub(r'\1', text)
    text = _RQ_DOUBLE_PERIOD.sub('。', text)
    text = re.sub(r'^[，,。；;：、]\s*', '', text)
    text = text.strip()

    # 9. If text is now just empty/punctuation-only, return empty
    if not text or all(c in '，,。；;：、 \t' for c in text):
        return ''

    return text


def set_cell_font(run, font_name='仿宋', font_size=Pt(12), bold=False):
    """Set font properties for a run."""
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold


def add_paragraph_with_inline(doc, text, style=None, font_name='仿宋',
                               font_size=Pt(12), bold=False, color=None,
                               alignment=None, space_after=Pt(6),
                               first_line_indent=None):
    """Add a paragraph with inline formatting (bold, code)."""
    para = doc.add_paragraph(style=style)

    if alignment is not None:
        para.alignment = alignment

    pf = para.paragraph_format
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = first_line_indent

    # Find all inline markers
    tokens = []
    for pattern, token_type in [(BOLD_PATTERN, 'bold'),
                                  (CODE_PATTERN, 'code')]:
        for m in pattern.finditer(text):
            tokens.append((m.start(), m.end(), m.group(1), token_type))

    tokens.sort(key=lambda x: x[0])

    # Build runs
    pos = 0
    for start, end, content, token_type in tokens:
        if pos < start:
            plain = text[pos:start]
            if plain:
                run = para.add_run(plain)
                set_cell_font(run, font_name, font_size, bold)
                if color:
                    run.font.color.rgb = color

        run = para.add_run(content)
        if token_type == 'bold':
            set_cell_font(run, '黑体', font_size, True)
        elif token_type == 'code':
            set_cell_font(run, 'Consolas', Pt(10.5), False)
        if color:
            run.font.color.rgb = color
        pos = end

    # Remaining plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        set_cell_font(run, font_name, font_size, bold)
        if color:
            run.font.color.rgb = color

    return para


def add_heading_para(doc, text, level):
    """Add a heading with proper Chinese formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            set_cell_font(run, '黑体', Pt(18), True)
        elif level == 2:
            set_cell_font(run, '黑体', Pt(15), True)
        elif level == 3:
            set_cell_font(run, '黑体', Pt(13), True)
    return heading


def add_requirement_annotation(doc, text):
    """Add a requirement annotation paragraph (blockquote style)."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.5)

    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F0E8')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

    pos = 0
    tokens = []
    for m in BOLD_PATTERN.finditer(text):
        tokens.append((m.start(), m.end(), m.group(1), 'bold'))

    for start, end, content, token_type in sorted(tokens):
        if pos < start:
            run = para.add_run(text[pos:start])
            set_cell_font(run, '楷体', Pt(11), False)
            run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        run = para.add_run(content)
        set_cell_font(run, '黑体', Pt(11), True)
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        pos = end

    if pos < len(text):
        run = para.add_run(text[pos:])
        set_cell_font(run, '楷体', Pt(11), False)
        run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

    return para


def add_image_with_caption(doc, image_path, caption_text):
    """Add an embedded image with a centered caption."""
    if os.path.exists(image_path):
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Cm(14.5))
        except Exception:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f'[图: {os.path.basename(image_path)}]')
            set_cell_font(run, '仿宋', Pt(10), False)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f'[图缺失: {os.path.basename(image_path)}]')
        set_cell_font(run, '仿宋', Pt(10), False)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = cap_para.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(12)

    # Parse bold markers in caption text
    pos = 0
    for m in BOLD_PATTERN.finditer(caption_text):
        if pos < m.start():
            run = cap_para.add_run(caption_text[pos:m.start()])
            set_cell_font(run, '楷体', Pt(9), False)
        run = cap_para.add_run(m.group(1))
        set_cell_font(run, '黑体', Pt(9), True)
        pos = m.end()

    if pos < len(caption_text):
        run = cap_para.add_run(caption_text[pos:])
        set_cell_font(run, '楷体', Pt(9), False)

    for run in cap_para.runs:
        if not run.bold:
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    return cap_para


def add_code_block(doc, code_lines):
    """Render a fenced code block as monospace paragraphs with gray background."""
    for line in code_lines:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(0.5)
        pPr = para._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F4F5F7')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)
        run = para.add_run(line if line else ' ')
        set_cell_font(run, 'Consolas', Pt(9), False)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_page_break(doc):
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def set_table_cell_font(cell, text, font_name='仿宋', font_size=Pt(10), bold=False):
    """Set text and font in a table cell. Clears existing paragraphs first."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if bold:
        run = p.add_run(text)
        set_cell_font(run, '黑体', font_size, True)
    else:
        run = p.add_run(text)
        set_cell_font(run, font_name, font_size, False)


def add_table(doc, table_lines):
    """Convert markdown pipe-table lines to a Word table."""
    # Parse rows
    rows = []
    is_separator = re.compile(r'^\|[\s\-:|]+\|$')
    for line in table_lines:
        if is_separator.match(line):
            continue  # Skip separator row
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    num_rows = len(rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = table.cell(ri, ci)
            is_header = (ri == 0)
            set_table_cell_font(cell, cell_text,
                                font_name='黑体' if is_header else '仿宋',
                                font_size=Pt(10) if is_header else Pt(10),
                                bold=is_header)
            # Header background
            if is_header:
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '1a3a5c')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
                # White text for header
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Add space after table
    doc.add_paragraph()


# Chapter outline_id → Chinese chapter name mapping
CHAPTER_NAME_MAP = {
    'CH-01': '第一章 项目理解与总体技术方案',
    'CH-02': '第二章 资质、保密与合规保障',
    'CH-03': '第三章 总体集成通用技术方案',
    'CH-04': '第四章 人机交互界面设计',
    'CH-05': '第五章 数据库模块技术方案',
    'CH-06': '第六章 材料设计模块技术方案',
    'CH-07': '第七章 热环境确定模块技术方案',
    'CH-08': '第八章 热环境响应评估模块技术方案',
    'CH-09': '第九章 运行环境兼容性与离线部署方案',
    'CH-10': '第十章 集成对接与数据贯通方案',
    'CH-11': '第十一章 项目实施与交付方案',
    'CH-12': '第十二章 测试、验收与质量保障',
    'CH-13': '第十三章 售后服务与质保承诺',
}


def add_bidding_comparison_table(doc, trace_path, chapter_name_map=None):
    """Generate and append the bidding requirements vs technical response comparison table."""
    if chapter_name_map is None:
        chapter_name_map = {}
    add_page_break(doc)
    add_heading_para(doc, '技术响应与投标要求对照表', 1)

    add_paragraph_with_inline(doc,
        '以下对照表逐条梳理招标文件中的技术要求（含★实质性条款、强制性条款及一般性条款），'
        '与投标文件技术方案各章节的响应位置建立一一对应关系，供评标专家逐条核实。'
        '所有条款均已在本投标文件中完整响应，无偏离、无遗漏。',
        font_name='仿宋', font_size=Pt(12), first_line_indent=Cm(0.74))

    # Load requirements from trace.json
    requirements = []
    if trace_path and trace_path.exists():
        try:
            with open(trace_path, 'r', encoding='utf-8') as f:
                trace = json.load(f)
            for req in trace.get('requirements', []):
                req_id = req.get('req_id', '')
                req_text = req.get('req_text', req.get('text', ''))
                priority = req.get('priority', '')
                source = req.get('source', req.get('section', ''))
                covered_by = req.get('covered_by', [])
                outline_id = covered_by[0] if covered_by else req.get('outline_id', '')
                chapter_name = chapter_name_map.get(outline_id, outline_id)
                body_ids = covered_by[1:] if len(covered_by) > 1 else req.get('body_ids', [])
                body_ref = '、'.join(body_ids[:6]) + ('…' if len(body_ids) > 6 else '')
                requirements.append({
                    'id': req_id,
                    'text': req_text,
                    'priority': priority,
                    'source': source,
                    'chapter': chapter_name,
                    'sections': body_ref,
                })
        except Exception as e:
            print(f"  WARNING: Failed to load trace.json: {e}")

    if not requirements:
        add_paragraph_with_inline(doc,
            '对照表数据来源 trace.json 不可用，请确保已运行 s1 阶段招标文件解析。',
            font_name='仿宋', font_size=Pt(12), first_line_indent=Cm(0.74))
        return

    # Summary paragraph
    star_count = sum(1 for r in requirements if '★' in r['priority'])
    mandatory_count = sum(1 for r in requirements if '强制' in r['priority'])
    general_count = sum(1 for r in requirements if '一般' in r['priority'])
    add_paragraph_with_inline(doc,
        f'本次投标共响应招标技术要求 {len(requirements)} 条，'
        f'其中★实质性条款 {star_count} 条、强制性条款 {mandatory_count} 条、'
        f'一般性条款 {general_count} 条。所有条款均已在技术方案相应章节中完整响应，无偏离。',
        font_name='仿宋', font_size=Pt(12), first_line_indent=Cm(0.74))

    # Build comparison table — 6 columns
    num_cols = 6
    table = doc.add_table(rows=len(requirements) + 1, cols=num_cols)
    table.style = 'Table Grid'

    headers = ['序号', '条款编号', '条款性质', '技术要求原文', '响应章节', '响应要点']
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        set_table_cell_font(cell, h, font_name='黑体', font_size=Pt(8), bold=True)
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1a3a5c')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, req in enumerate(requirements):
        # 序号
        set_table_cell_font(table.cell(ri + 1, 0), str(ri + 1), font_size=Pt(8))

        # 条款编号
        set_table_cell_font(table.cell(ri + 1, 1), req['id'],
                            font_name='Consolas', font_size=Pt(7), bold=True)

        # 条款性质 — with color coding
        priority = req['priority']
        type_cell = table.cell(ri + 1, 2)
        if '★' in priority or '实质性' in priority:
            set_table_cell_font(type_cell, priority, font_size=Pt(8), bold=True)
            for p in type_cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif '强制' in priority:
            set_table_cell_font(type_cell, priority, font_size=Pt(8), bold=True)
            for p in type_cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xD9, 0x75, 0x00)
        else:
            set_table_cell_font(type_cell, priority, font_size=Pt(8))

        # 技术要求原文
        set_table_cell_font(table.cell(ri + 1, 3), req['text'][:200], font_size=Pt(7))

        # 响应章节
        set_table_cell_font(table.cell(ri + 1, 4), req['chapter'], font_size=Pt(7))

        # 响应要点 — section references
        set_table_cell_font(table.cell(ri + 1, 5), req['sections'], font_size=Pt(7))

    # Set column widths
    widths = [Cm(0.8), Cm(2.0), Cm(1.5), Cm(5.5), Cm(3.0), Cm(3.2)]
    for row in table.rows:
        for ci, width in enumerate(widths):
            row.cells[ci].width = width

    doc.add_paragraph()


def export_to_docx(combined_md=None, s6_dir=None, s8_dir=None,
                   trace_path=None, chapter_name_map=None, project_name=None):
    """Main export function.

    Args:
        combined_md: Path to s6_combined_bid_document.md
        s6_dir: Path to s6_synthesis directory (for resolving image paths)
        s8_dir: Path to s8_final output directory
        trace_path: Path to trace.json
        chapter_name_map: Dict mapping outline_id to chapter name
        project_name: Project name for output filename
    """
    output_dir = Path.cwd() / 'output'
    combined_md = Path(combined_md) if combined_md else output_dir / 's6_synthesis' / 's6_combined_bid_document.md'
    s6_dir = Path(s6_dir) if s6_dir else output_dir / 's6_synthesis'
    s8_dir = Path(s8_dir) if s8_dir else output_dir / 's8_final'
    trace_path = Path(trace_path) if trace_path else output_dir / 'trace.json'
    if chapter_name_map is None:
        chapter_name_map = {}
    if project_name is None:
        project_name = '投标项目'

    print("Reading combined markdown...")
    with open(combined_md, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = Document()
    doc.settings.element.append(OxmlElement('w:doNotCompressPictures'))

    # ─── Set up styles ───
    style = doc.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ─── Process lines ───
    i = 0
    image_count = 0
    total_lines = len(lines)

    while i < total_lines:
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Fenced code block
        fence_match = FENCE_PATTERN.match(line.strip())
        if fence_match:
            code_lines = []
            i += 1
            while i < total_lines:
                if FENCE_PATTERN.match(lines[i].strip()):
                    i += 1
                    break
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                add_code_block(doc, code_lines)
            continue

        # Heading
        heading_match = HEADING_PATTERN.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            add_heading_para(doc, text, level)
            if level == 1 and i > 10:
                add_page_break(doc)
            i += 1
            continue

        # Horizontal rule
        if HR_PATTERN.match(line.strip()):
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote — skip requirement annotations in body text
        # (requirements are covered in the appendix bidding comparison table)
        blockquote_match = BLOCKQUOTE_PATTERN.match(line)
        if blockquote_match:
            text = blockquote_match.group(1)
            if text.strip().startswith('**【对标条款'):
                i += 1
                continue  # skip inline requirement annotations
            text = clean_body_text(text)
            if not text:
                i += 1
                continue
            add_requirement_annotation(doc, text)
            i += 1
            continue

        # Image — with blank-line-tolerant caption detection
        image_match = IMAGE_PATTERN.match(line.strip())
        if image_match:
            alt_text = image_match.group(1)
            img_rel_path = image_match.group(2)
            img_abs_path = str(s6_dir / img_rel_path)

            caption = alt_text  # fallback
            i += 1  # skip image line

            # Skip blank lines between image and caption
            while i < total_lines and not lines[i].strip():
                i += 1

            # Check if next non-blank line is a caption
            if i < total_lines:
                next_line = lines[i].strip()
                caption_match = CAPTION_PATTERN.match(next_line)
                if caption_match:
                    caption = caption_match.group(1)
                    i += 1  # skip caption line

            add_image_with_caption(doc, img_abs_path, caption)
            image_count += 1
            continue

        # Orphaned caption line (shouldn't happen after fix, but handle gracefully)
        stripped = line.strip()
        caption_match = CAPTION_PATTERN.match(stripped)
        # Skip pure bold lines (**text**) — CAPTION_PATTERN greedily matches them
        if caption_match and not (stripped.startswith('**') and stripped.endswith('**')):
            caption_text = caption_match.group(1)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Render with bold parsing
            pos = 0
            for m in BOLD_PATTERN.finditer(caption_text):
                if pos < m.start():
                    run = para.add_run(caption_text[pos:m.start()])
                    set_cell_font(run, '楷体', Pt(9), False)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run = para.add_run(m.group(1))
                set_cell_font(run, '黑体', Pt(9), True)
                pos = m.end()
            if pos < len(caption_text):
                run = para.add_run(caption_text[pos:])
                set_cell_font(run, '楷体', Pt(9), False)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Table detection: consecutive lines starting and ending with |
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines = []
            while i < total_lines and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if table_lines:
                add_table(doc, table_lines)
            continue

        # List items
        list_match = LIST_ITEM_PATTERN.match(line)
        if list_match:
            text = clean_body_text(list_match.group(2))
            if not text:
                i += 1
                continue
            add_paragraph_with_inline(doc, f'• {text}', font_name='仿宋',
                                       font_size=Pt(12))
            last_para = doc.paragraphs[-1]
            last_para.paragraph_format.left_indent = Cm(1.0)
            last_para.paragraph_format.first_line_indent = Cm(-0.5)
            last_para.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Regular paragraph
        line_text = clean_body_text(line.strip())
        if not line_text:
            i += 1
            continue
        add_paragraph_with_inline(doc, line_text, font_name='仿宋',
                                   font_size=Pt(12),
                                   first_line_indent=Cm(0.74))
        i += 1

    # ─── Append bidding requirements comparison table ───
    add_bidding_comparison_table(doc, trace_path, chapter_name_map)

    # ─── Save ───
    import datetime
    s8_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_name = project_name.replace(' ', '_').replace('/', '_')
    output_path = str(s8_dir / f'投标文件_{safe_name}_技术方案_{ts}.docx')
    doc.save(output_path)

    print(f"\n{'='*60}")
    print(f"Stage 8 Export Complete")
    print(f"{'='*60}")
    print(f"  Output:      {output_path}")
    print(f"  Images embedded: {image_count}")
    print(f"  Paragraphs:  {len(doc.paragraphs)}")

    return output_path


if __name__ == '__main__':
    export_to_docx()
